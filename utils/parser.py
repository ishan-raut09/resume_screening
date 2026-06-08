import os
import pdfplumber
import PyPDF2
import docx
import re

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file using pdfplumber, with a fallback to PyPDF2."""
    text = ""
    try:
        # Try pdfplumber first as it is generally more accurate for layout
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"pdfplumber failed for {pdf_path}: {e}. Trying PyPDF2 fallback...")
        try:
            with open(pdf_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e2:
            print(f"PyPDF2 fallback also failed for {pdf_path}: {e2}")
    
    return text.strip()

def extract_text_from_docx(docx_path):
    """Extract text from a DOCX file using python-docx."""
    text = ""
    try:
        doc = docx.Document(docx_path)
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
        # Also extract text from tables within docx
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"python-docx failed for {docx_path}: {e}")
    return text.strip()

def extract_text(file_path):
    """Extract text from a file based on its extension."""
    _, ext = os.path.splitext(file_path.lower())
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

def extract_email(text):
    email = re.search(r'[\w\.-]+@[\w\.-]+', text)
    return email.group(0) if email else "Not Found"

def extract_phone(text):
    phone = re.search(r'\+?\d[\d -]{8,12}\d', text)
    return phone.group(0).strip() if phone else "Not Found"

def extract_name(text):
    # Simple heuristic to extract name: often the first or second line in a resume
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    for line in lines[:5]: # look at first 5 non-empty lines
        if 1 <= len(line.split()) <= 4 and re.match(r'^[A-Za-z\s\.-]+$', line):
            return line.title()
    return "Not Found"
